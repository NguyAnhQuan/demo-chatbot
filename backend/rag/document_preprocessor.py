from .text_processor import TextProcessor
import os
import re
import unicodedata
import fitz  # PyMuPDF
import easyocr
from docx import Document
import win32com.client as win32

class DocumentPreprocessor(TextProcessor):
    """Handles text preprocessing and optional OCR."""
    def __init__(self, use_ocr: bool = False, ocr_model: str = 'vie+eng'): 
        self.use_ocr = use_ocr
        self.ocr_lang = ocr_model
        self._ocr_reader = None  # Cache OCR reader

    def process(self, file_path: str) -> str:
        text = self._extract_text(file_path)
        cleaned_text = self._preprocess_text(text)
        return cleaned_text
    
    def _extract_text(self, file_path: str) -> str:
        """Extract text from various file formats with optional OCR."""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_ext = os.path.splitext(file_path)[1].lower()
        
        try:
            # PDF files
            if file_ext == '.pdf':
                return self._extract_from_pdf(file_path)
            
            # docx documents
            elif file_ext == '.docx':
                return self._extract_from_docx(file_path)

            else:
                raise ValueError(f"Unsupported file format: {file_ext}")
                
        except Exception as e:
            raise Exception(f"Error extracting text from {file_path}: {str(e)}")
        

    def _extract_from_pdf(self, file_path: str) -> str:
    #Extract text from PDF, with OCR fallback for scanned PDFs.
        text = ""
        doc = fitz.open(file_path)
        
        for page_num in range(doc.page_count):
            page = doc[page_num]
            page_text = page.get_text()
            
            # If no text found and OCR is enabled, try OCR
            if not page_text.strip() and self.use_ocr:
                # Convert page to image and apply OCR
                pix = page.get_pixmap()
                img_data = pix.tobytes("png")
                
                # Use EasyOCR for Vietnamese text
                reader = easyocr.Reader([self.ocr_lang.split('+')[0], 'en'])
                results = reader.readtext(img_data)
                page_text = ' '.join([result[1] for result in results])
            
            text += page_text + "\n"
        
        doc.close()
        return text
    
    def _extract_from_docx(self, file_path: str) -> str:
        """Extract text from Word documents."""
        doc = Document(file_path)
        text = ""
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
            text += "\n"
        
        return text
    

    def _preprocess_text(self, text: str) -> str:
        """Preprocess text by removing special characters, extra whitespace, and normalizing unicode."""
        if not text:
            return ""
        
        # 1. Normalize Unicode (important for Vietnamese text)
        text = unicodedata.normalize('NFC', text)
        
        # 2. Remove or replace problematic characters
        # Remove zero-width characters
        text = re.sub(r'[\u200b\u200c\u200d\ufeff]', '', text)
        
        # Replace different types of quotes with standard ones
        text = re.sub(r'[""''`´]', '"', text)
        text = re.sub(r'[''‚]', "'", text)
        
        # Replace different types of dashes with standard dash
        text = re.sub(r'[–—−]', '-', text)
        
        # Replace different types of ellipsis
        text = re.sub(r'[…]', '...', text)
        
        # 3. Clean up whitespace
        # Replace multiple spaces with single space
        text = re.sub(r' +', ' ', text)
        
        # Replace multiple newlines with double newline (paragraph break)
        text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)
        
        # Remove trailing whitespace from each line
        text = '\n'.join(line.rstrip() for line in text.split('\n'))
        
        # 4. Remove or clean OCR artifacts
        # Remove isolated single characters (common OCR errors)
        text = re.sub(r'\b[a-zA-Z]\b', '', text)
        
        # Remove excessive punctuation
        text = re.sub(r'[.]{4,}', '...', text)
        text = re.sub(r'[!]{2,}', '!', text)
        text = re.sub(r'[?]{2,}', '?', text)
        
        # 5. Clean up page markers and headers/footers
        # Remove page numbers at start of line
        text = re.sub(r'^\s*\d+\s*$', '', text, flags=re.MULTILINE)
        
        # Remove common header/footer patterns
        text = re.sub(r'--- Page \d+ ---\n?', '', text)
        
        # 6. Remove URLs and email addresses if not needed
        text = re.sub(r'http[s]?://(?:[a-zA-Z]|[0-9]|[$-_@.&+]|[!*\\(\\),]|(?:%[0-9a-fA-F][0-9a-fA-F]))+', '', text)
        text = re.sub(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', '', text)
        
        # 7. Remove excessive special characters but keep Vietnamese punctuation
        # Keep: . , ; : ! ? ( ) [ ] { } " ' - / \ và các ký tự tiếng Việt
        text = re.sub(r'[^\w\sàáãạảăắằẳẵặâấầẩẫậèéẹẻẽêềếểễệđìíĩỉịòóõọỏôốồổỗộơớờởỡợùúũụủưứừửữựỳýỵỷỹ.,:;!?()\[\]{}"\'\\/-]', ' ', text, flags=re.IGNORECASE)
        
        # 8. Final cleanup
        # Remove multiple consecutive spaces again
        text = re.sub(r' +', ' ', text)
        
        # Remove empty lines
        lines = [line.strip() for line in text.split('\n') if line.strip()]
        text = '\n'.join(lines)
        
        # Trim leading and trailing whitespace
        text = text.strip()
        
        return text
